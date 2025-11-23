'''
Business: Creates default DOCX template with all required variables
Args: event - dict with httpMethod, context - object with request_id
Returns: Base64 encoded DOCX template file
'''

import json
import os
import base64
from typing import Dict, Any
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import io

def handler(event: Dict[str, Any], context: Any) -> Dict[str, Any]:
    method: str = event.get('httpMethod', 'GET')
    
    if method == 'OPTIONS':
        return {
            'statusCode': 200,
            'headers': {
                'Access-Control-Allow-Origin': '*',
                'Access-Control-Allow-Methods': 'GET, POST, OPTIONS',
                'Access-Control-Allow-Headers': 'Content-Type',
                'Access-Control-Max-Age': '86400'
            },
            'body': ''
        }
    
    if method not in ['GET', 'POST']:
        return {
            'statusCode': 405,
            'headers': {'Content-Type': 'application/json', 'Access-Control-Allow-Origin': '*'},
            'isBase64Encoded': False,
            'body': json.dumps({'error': 'Method not allowed'})
        }
    
    doc = Document()
    
    section = doc.sections[0]
    section.page_height = Inches(11.69)
    section.page_width = Inches(8.27)
    section.top_margin = Inches(0.79)
    section.bottom_margin = Inches(0.79)
    section.left_margin = Inches(1.18)
    section.right_margin = Inches(0.59)
    
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run('ПРОТОКОЛ')
    run.bold = True
    run.font.size = Pt(14)
    
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run('заседания военно-врачебной комиссии')
    run.font.size = Pt(12)
    
    date_para = doc.add_paragraph()
    date_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    date_para.add_run('{date}')
    
    doc.add_paragraph()
    
    meeting_para = doc.add_paragraph()
    meeting_para.add_run('Заседание № {meetingNumber}')
    
    doc.add_paragraph()
    doc.add_paragraph('Рассмотрено протоколов: {protocolCount}')
    doc.add_paragraph()
    
    doc.add_paragraph('─' * 80)
    doc.add_paragraph()
    
    loop_start = doc.add_paragraph()
    loop_start.add_run('{#protocols}')
    
    protocol_header = doc.add_paragraph()
    protocol_header.add_run('ПРОТОКОЛ № {number}')
    protocol_header.runs[0].bold = True
    
    doc.add_paragraph()
    
    doc.add_paragraph('ФИО: {fio}')
    doc.add_paragraph('Дата рождения: {birthDate}')
    doc.add_paragraph('Воинское звание: {rank}')
    doc.add_paragraph('Должность: {position}')
    doc.add_paragraph('Воинская часть: {militaryUnit}')
    
    doc.add_paragraph()
    
    contract_block = doc.add_paragraph()
    contract_block.add_run('Условия службы: По контракту')
    doc.add_paragraph('  - Дата контракта: {contractDate}')
    doc.add_paragraph('  - Контракт подписан: {contractSigner}')
    
    doc.add_paragraph()
    
    mobilization_block = doc.add_paragraph()
    mobilization_block.add_run('Условия службы: По мобилизации')
    doc.add_paragraph('  - Дата мобилизации: {mobilizationDate}')
    doc.add_paragraph('  - Мобилизован из: {mobilizationSource}')
    
    doc.add_paragraph()
    doc.add_paragraph('─' * 80)
    doc.add_paragraph()
    
    loop_end = doc.add_paragraph()
    loop_end.add_run('{/protocols}')
    
    doc.add_paragraph()
    doc.add_paragraph()
    
    signature_para = doc.add_paragraph()
    signature_para.add_run('Председатель комиссии: _________________')
    
    doc.add_paragraph()
    doc.add_paragraph('Члены комиссии:')
    doc.add_paragraph('    _________________')
    doc.add_paragraph('    _________________')
    
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    
    docx_bytes = buffer.read()
    docx_base64 = base64.b64encode(docx_bytes).decode('utf-8')
    
    return {
        'statusCode': 200,
        'headers': {
            'Content-Type': 'application/json',
            'Access-Control-Allow-Origin': '*'
        },
        'isBase64Encoded': False,
        'body': json.dumps({
            'fileContent': docx_base64,
            'fileName': 'template_zasedanie.docx',
            'message': 'Default template created successfully'
        })
    }
